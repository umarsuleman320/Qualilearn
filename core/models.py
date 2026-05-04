from django.db import models
from django.contrib.auth.models import User
from django.utils import timezone
from django.db.models.signals import post_save
from django.dispatch import receiver
import os

class LearningCategory(models.Model):
    name = models.CharField(max_length=100)
    icon_class = models.CharField(max_length=50, default='bi-book', help_text="Bootstrap icon class")
    order = models.IntegerField(default=0)

    class Meta:
        verbose_name_plural = "Learning Categories"
        ordering = ['order', 'name']

    def __str__(self):
        return self.name

class Subject(models.Model):
    category = models.ForeignKey(LearningCategory, on_delete=models.SET_NULL, null=True, blank=True, related_name='subjects')
    name = models.CharField(max_length=100)
    document = models.FileField(upload_to='question_banks/', null=True, blank=True, help_text="Upload a .docx file containing the questions. Existing questions will be wiped and replaced.")
    flashcard_document = models.FileField(upload_to='flashcard_banks/', null=True, blank=True, help_text="Upload a .docx file for Flashcards (Format: Front: text \\n Back: text). Existing flashcards will be wiped and replaced.")

    def __str__(self):
        return self.name

class Question(models.Model):
    subject = models.ForeignKey(Subject, on_delete=models.CASCADE)
    text = models.TextField()
    option_a = models.CharField(max_length=300)
    option_b = models.CharField(max_length=300)
    option_c = models.CharField(max_length=300)
    option_d = models.CharField(max_length=300)
    correct_option = models.CharField(max_length=1)   # 'a', 'b', 'c', or 'd'
    explanation = models.TextField(blank=True)

    def __str__(self):
        return self.text[:80]

class Flashcard(models.Model):
    subject = models.ForeignKey(Subject, on_delete=models.CASCADE)
    front = models.TextField()
    back = models.TextField()

    def __str__(self):
        return f"Flashcard: {self.front[:50]}"

class UserProgress(models.Model):
    user = models.ForeignKey(User, on_delete=models.CASCADE)
    subject = models.ForeignKey(Subject, on_delete=models.CASCADE)
    score = models.IntegerField(default=0)
    total_questions = models.IntegerField(default=15)
    completed = models.BooleanField(default=False)
    last_attempt = models.DateTimeField(default=timezone.now)

    def percentage(self):
        if self.total_questions == 0:
            return 0
        return round((self.score / self.total_questions) * 100)

    def __str__(self):
        return f"{self.user.username} - {self.subject.name} ({self.percentage()}%)"

class UserProfile(models.Model):
    user = models.OneToOneField(User, on_delete=models.CASCADE)
    streak_count = models.IntegerField(default=0)
    total_study_minutes = models.IntegerField(default=0)
    last_active_date = models.DateField(null=True, blank=True)
    profile_picture = models.ImageField(upload_to='profile_pics/', blank=True, null=True)

    def __str__(self):
        return f"{self.user.username}'s Profile"

class VocationalSkill(models.Model):
    CATEGORY_CHOICES = [
        ('Practical Skills', 'Practical Skills'),
        ('Soft / Digital Skills', 'Soft / Digital Skills'),
    ]
    name = models.CharField(max_length=150)
    category = models.CharField(max_length=50, choices=CATEGORY_CHOICES, default='Practical Skills')
    youtube_url = models.URLField(max_length=500)
    icon_class = models.CharField(max_length=50, blank=True, help_text='Bootstrap icon class (e.g., bi-tools)')
    
    def __str__(self):
        return f"{self.name} ({self.category})"

# Signals to automatically create/update UserProfile
@receiver(post_save, sender=User)
def create_user_profile(sender, instance, created, **kwargs):
    if created:
        UserProfile.objects.get_or_create(user=instance)

@receiver(post_save, sender=User)
def save_user_profile(sender, instance, **kwargs):
    try:
        instance.userprofile.save()
    except Exception:
        pass

# Signal to parse the DOCX file whenever a Subject is saved with a document
@receiver(post_save, sender=Subject)
def parse_and_sync_docx(sender, instance, **kwargs):
    if instance.document and os.path.exists(instance.document.path):
        from docx import Document
        
        # Clear existing questions for this subject to prevent duplicates
        Question.objects.filter(subject=instance).delete()
        
        doc = Document(instance.document.path)
        
        # A list to collect our blocks
        blocks = []
        current_block = {}
        
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            
            lower_text = text.lower()
            
            if lower_text.startswith('question:'):
                if current_block: blocks.append(current_block)
                current_block = {'text': text.split(':', 1)[1].strip(), 'options': {}, 'answer': None}
            elif current_block:
                # Check for options A, B, C, D with various delimiters
                import re
                option_match = re.match(r'^([A-D])[\:\)\.\-\s]', text, re.I)
                if option_match:
                    opt_letter = option_match.group(1).upper()
                    # Strip the "A:" part
                    opt_content = re.sub(r'^[A-D][\:\)\.\-\s]\s*', '', text, flags=re.I).strip()
                    current_block['options'][opt_letter] = opt_content
                elif lower_text.startswith('answer:'):
                    ans_val = text.split(':', 1)[1].strip().lower()
                    if ans_val:
                        current_block['answer'] = ans_val[0]
        
        # Don't forget the last one
        if current_block: blocks.append(current_block)
        
        for b in blocks:
            if b.get('text') and b.get('answer'):
                Question.objects.create(
                    subject=instance,
                    text=b['text'],
                    option_a=b['options'].get('A', ''),
                    option_b=b['options'].get('B', ''),
                    option_c=b['options'].get('C', ''),
                    option_d=b['options'].get('D', ''),
                    correct_option=b['answer']
                )

@receiver(post_save, sender=Subject)
def parse_and_sync_flashcard_docx(sender, instance, **kwargs):
    if instance.flashcard_document and os.path.exists(instance.flashcard_document.path):
        from docx import Document
        
        # Clear existing flashcards for this subject to prevent duplicates
        Flashcard.objects.filter(subject=instance).delete()
        
        try:
            doc = Document(instance.flashcard_document.path)
            current_flashcard = None
            
            for p in doc.paragraphs:
                text = p.text.strip()
                if not text:
                    continue
                    
                lower_text = text.lower()
                
                # Start of a new flashcard
                if lower_text.startswith('front:'):
                    current_flashcard = Flashcard(subject=instance)
                    parts = text.split(":", 1)
                    if len(parts) > 1:
                        current_flashcard.front = parts[1].strip()
                
                elif current_flashcard and lower_text.startswith('back:'):
                    parts = text.split(":", 1)
                    if len(parts) > 1:
                        current_flashcard.back = parts[1].strip()
                    
                    if current_flashcard.front and current_flashcard.back:
                        current_flashcard.save()
                    current_flashcard = None
                    
        except Exception as e:
            print("Flashcard document parsing failed:", e)

class BroadcastMessage(models.Model):
    subject = models.CharField(max_length=200)
    message = models.TextField()
    created_at = models.DateTimeField(auto_now_add=True)
    sent = models.BooleanField(default=False)

    def __str__(self):
        return self.subject

class LearningTopic(models.Model):
    subject = models.ForeignKey(Subject, on_delete=models.CASCADE, related_name='learning_topics', null=True)
    name = models.CharField(max_length=200)
    youtube_url = models.URLField(max_length=500, blank=True, null=True)
    order = models.IntegerField(default=0)

    class Meta:
        ordering = ['order', 'name']

    def __str__(self):
        return f"{self.subject.name if self.subject else 'No Subject'} - {self.name}"

class LearningProgress(models.Model):
    user = models.ForeignKey(User, on_delete=models.CASCADE)
    topic = models.ForeignKey(LearningTopic, on_delete=models.CASCADE)
    is_completed = models.BooleanField(default=False)
    completed_at = models.DateTimeField(null=True, blank=True)

    class Meta:
        unique_together = ('user', 'topic')

    def __str__(self):
        status = "Completed" if self.is_completed else "In Progress"
        return f"{self.user.username} - {self.topic.name} ({status})"
